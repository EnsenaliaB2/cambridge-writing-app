import os
from typing import List, Dict, Any, Tuple

from docx import Document
from docx.shared import RGBColor, Inches
from docx2pdf import convert

from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


NAVY = RGBColor(20, 35, 90)
SOFT_RED = RGBColor(200, 90, 90)


def _set_run(run, color: RGBColor, underline: bool = False):
    run.font.color.rgb = color
    run.underline = underline


def _set_paragraph_color(paragraph, color: RGBColor):
    for r in paragraph.runs:
        _set_run(r, color=color, underline=False)


def _clear_paragraph(paragraph):
    paragraph.text = ""


def _find_paragraph_index_contains(doc: Document, needle: str) -> int:
    for i, p in enumerate(doc.paragraphs):
        if needle in (p.text or ""):
            return i
    return -1


def _insert_paragraph_after_paragraph(base_para: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    base_para._p.addnext(new_p)
    new_para = Paragraph(new_p, base_para._parent)
    if text:
        run = new_para.add_run(text)
        _set_run(run, color=NAVY)
    return new_para


def _insert_paragraph_after_index(doc: Document, paragraph_index: int, text: str = "") -> Paragraph:
    return _insert_paragraph_after_paragraph(doc.paragraphs[paragraph_index], text)


# ----------------------------
# Student answer underline
# ----------------------------
def _write_inline_underlines(paragraph: Paragraph, full_text: str, corrections: List[Dict[str, Any]]):
    _clear_paragraph(paragraph)

    fragments = []
    for c in corrections or []:
        if isinstance(c, dict):
            frag = (c.get("error_fragment") or "").strip()
            if frag:
                fragments.append(frag)

    idx = 0
    used = set()

    while idx < len(full_text):
        next_pos = None
        next_frag = None

        for frag in fragments:
            if frag in used:
                continue
            pos = full_text.find(frag, idx)
            if pos != -1 and (next_pos is None or pos < next_pos):
                next_pos = pos
                next_frag = frag

        if next_pos is None:
            r = paragraph.add_run(full_text[idx:])
            _set_run(r, NAVY)
            break

        if next_pos > idx:
            r = paragraph.add_run(full_text[idx:next_pos])
            _set_run(r, NAVY)

        r_err = paragraph.add_run(next_frag)
        _set_run(r_err, SOFT_RED, underline=True)

        used.add(next_frag)
        idx = next_pos + len(next_frag)


def _write_suggested_corrections_after_paragraph(anchor_para: Paragraph, corrections: List[Dict[str, Any]]):
    if not corrections:
        return

    p_title = _insert_paragraph_after_paragraph(anchor_para, "Suggested corrections")
    _set_paragraph_color(p_title, NAVY)

    last_para = p_title
    for c in corrections[:8]:
        if not isinstance(c, dict):
            continue

        err = (c.get("error_fragment") or "").strip()
        sug = (c.get("suggestion") or "").strip()

        p = _insert_paragraph_after_paragraph(last_para, "")
        last_para = p

        if err:
            r1 = p.add_run(err)
            _set_run(r1, SOFT_RED, underline=True)
        if sug:
            r2 = p.add_run(" → " + sug)
            _set_run(r2, NAVY)


# ----------------------------
# Score table
# ----------------------------
def _fill_score_table(doc: Document, score: Dict[str, Any]):
    target = None
    for t in doc.tables:
        try:
            header = " ".join(cell.text.strip() for cell in t.rows[0].cells)
            if ("Criterion" in header) and ("Score" in header) and ("Max" in header):
                target = t
                break
        except Exception:
            continue

    if not target:
        return

    mapping = {
        "Content": score.get("content", ""),
        "Communicative Achievement": score.get("communicative_achievement", ""),
        "Organisation": score.get("organisation", ""),
        "Language": score.get("language", ""),
    }

    for row in target.rows[1:]:
        crit = row.cells[0].text.strip()
        if crit in mapping:
            row.cells[1].text = str(mapping[crit])

        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run(r, NAVY)

    band = score.get("overall_band", "B2")
    total = score.get("total", "")

    idx_band = _find_paragraph_index_contains(doc, "Overall Band:")
    if idx_band != -1:
        doc.paragraphs[idx_band].text = f"Overall Band: {band}"
        _set_paragraph_color(doc.paragraphs[idx_band], NAVY)

        if idx_band + 1 < len(doc.paragraphs) and "Total:" in (doc.paragraphs[idx_band + 1].text or ""):
            doc.paragraphs[idx_band + 1].text = f"Total: {total} / 20"
            _set_paragraph_color(doc.paragraphs[idx_band + 1], NAVY)


# ----------------------------
# Detailed section 3 feedback
# ----------------------------
def _insert_section3_feedback(doc: Document, data: Dict[str, Any]):
    idx_section3 = _find_paragraph_index_contains(doc, "3. Feedback and Score")
    if idx_section3 == -1:
        return

    last_para = doc.paragraphs[idx_section3]

    overall_comment = (data.get("overall_comment") or "").strip()
    if overall_comment:
        p = _insert_paragraph_after_paragraph(last_para, overall_comment)
        _set_paragraph_color(p, NAVY)
        last_para = p

    criterion_feedback = data.get("criterion_feedback", {}) or {}
    criterion_evidence = data.get("criterion_evidence", {}) or {}

    sections = [
        ("1. Content", criterion_feedback.get("content", ""), criterion_evidence.get("content", [])),
        ("2. Communicative Achievement", criterion_feedback.get("communicative_achievement", ""), criterion_evidence.get("communicative_achievement", [])),
        ("3. Organisation", criterion_feedback.get("organisation", ""), criterion_evidence.get("organisation", [])),
        ("4. Language", criterion_feedback.get("language", ""), criterion_evidence.get("language", [])),
    ]

    for title, body, evidence_list in sections:
        if title:
            p_title = _insert_paragraph_after_paragraph(last_para, title)
            _set_paragraph_color(p_title, NAVY)
            last_para = p_title

        if body:
            for line in str(body).splitlines():
                p_body = _insert_paragraph_after_paragraph(last_para, line)
                _set_paragraph_color(p_body, NAVY)
                last_para = p_body

        evidence_list = evidence_list or []
        evidence_list = [str(x).strip() for x in evidence_list if str(x).strip()]
        if evidence_list:
            p_ev_title = _insert_paragraph_after_paragraph(last_para, "Evidence from script:")
            _set_paragraph_color(p_ev_title, NAVY)
            last_para = p_ev_title

            for ev in evidence_list[:2]:
                p_ev = _insert_paragraph_after_paragraph(last_para, f"• {ev}")
                _set_paragraph_color(p_ev, NAVY)
                last_para = p_ev


# ----------------------------
# Bullets
# ----------------------------
def _fill_bullets_under_heading(doc: Document, heading_text: str, lines: List[str], max_lines: int = 10):
    idx = _find_paragraph_index_contains(doc, heading_text)
    if idx == -1:
        return

    if not lines:
        lines = ["—"]

    last_para = doc.paragraphs[idx]
    for line in lines[:max_lines]:
        p = _insert_paragraph_after_paragraph(last_para, "• " + str(line))
        _set_paragraph_color(p, NAVY)
        last_para = p


# ----------------------------
# Build report
# ----------------------------
def build_report_from_template(
    template_path: str,
    output_docx: str,
    task: str,
    student_answer: str,
    data: Dict[str, Any],
):
    doc = Document(template_path)

    # 1. Task (image + extracted text)
    idx_task = _find_paragraph_index_contains(doc, "1. Task")
    if idx_task != -1:
        task_image_path = (data.get("task_image_path") or "").strip()
        last = doc.paragraphs[idx_task]

        if task_image_path and os.path.exists(task_image_path):
            p_img = _insert_paragraph_after_paragraph(last, "")
            try:
                p_img.add_run().add_picture(task_image_path, width=Inches(5.8))
            except Exception:
                pass
            last = p_img

        task_text = (data.get("task_extracted") or "").strip()
        if not task_text:
            task_text = task or ""

        if task_text.strip():
            p_task = _insert_paragraph_after_paragraph(last, task_text)
            _set_paragraph_color(p_task, NAVY)

    # 2. Student Answer
    idx_student = _find_paragraph_index_contains(doc, "2. Student Answer")
    if idx_student != -1:
        p_ans = _insert_paragraph_after_index(doc, idx_student, "")
        _write_inline_underlines(p_ans, student_answer, data.get("corrections", []))
        _write_suggested_corrections_after_paragraph(p_ans, data.get("corrections", []))

    # 3. Feedback and Score
    _insert_section3_feedback(doc, data)
    _fill_score_table(doc, data.get("score", {}))

    # 4/5/6
    strengths = []
    s = data.get("strengths", {}) or {}
    for k in ["content", "communicative_achievement", "organisation", "language"]:
        strengths += (s.get(k, []) or [])

    weaknesses = []
    w = data.get("weaknesses", {}) or {}
    for k in ["content", "communicative_achievement", "organisation", "language"]:
        weaknesses += (w.get(k, []) or [])

    improvements = data.get("improvement_summary", []) or []

    _fill_bullets_under_heading(doc, "4. Strengths", strengths)
    _fill_bullets_under_heading(doc, "5. Weaknesses", weaknesses)
    _fill_bullets_under_heading(doc, "6. Improvement Summary", improvements)

    # 7. Model Answer
    idx_model = _find_paragraph_index_contains(doc, "7. Model Answer")
    if idx_model != -1:
        p_model = _insert_paragraph_after_index(doc, idx_model, data.get("model_answer", "") or "")
        _set_paragraph_color(p_model, NAVY)

    doc.save(output_docx)


def create_report(data: Dict[str, Any], task: str, answer: str) -> Tuple[str, str]:
    template_path = "template.docx"
    out_docx = "writing_feedback.docx"
    out_pdf = "writing_feedback.pdf"

    if not os.path.exists(template_path):
        raise FileNotFoundError("template.docx not found")

    build_report_from_template(template_path, out_docx, task, answer, data)

    pdf_path = ""
    try:
        convert(out_docx, out_pdf)
        pdf_path = out_pdf
    except Exception:
        pass

    return out_docx, pdf_path